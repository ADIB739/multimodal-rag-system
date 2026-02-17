# uifinal_full.py
# Streamlit app: Fast RAG Q&A (WhisperX) + Map-Reduce Summarization + Exports + Flashcards
# Notes:
# - Uses your requested hard-coded keys (GROQ, HF, NEBIUS).
# - Summarization can run any time AFTER Transcribe+Diarize has been run once in this session.
# - "Prefer official YouTube transcript" is supported via youtube-transcript-api (no Google API key).
# - Exports conversation as Markdown / DOCX / PDF (DOCX requires `python-docx`, PDF requires `fpdf2`).
# - NEW: Microphone input for voice-based questions. Requires `streamlit-mic-recorder`, `soundfile`, `numpy`.

import os
import re
import gc
import time
import tempfile
import ffmpeg
import yt_dlp
import torch
import whisperx
import streamlit as st
from datetime import timedelta
from typing import Optional, List, Dict, Any, Tuple
from enum import Enum
from io import BytesIO
import librosa
from transformers import pipeline

# ---- RAG / LLM imports
from langchain_groq import ChatGroq
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_classic.chains.combine_documents import create_stuff_documents_chain

from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv
load_dotenv()
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ModuleNotFoundError:
    from langchain_community.embeddings import HuggingFaceEmbeddings  # fallback

# ---- Nebius (OpenAI-compatible) for summarization & flashcards
from openai import OpenAI

# ---- YouTube official transcript (no API key)
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
import os

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
HF_TOKEN = os.getenv("HF_TOKEN")
NEBIUS_API_KEY = os.getenv("NEBIUS_API_KEY")

# ---- Optional exports (docx/pdf)
try:
    from docx import Document as DOCXDocument
except Exception:
    DOCXDocument = None

try:
    from fpdf import FPDF
except Exception:
    FPDF = None

# Catch specific FPDF error class safely (works across fpdf versions)
try:
    from fpdf.errors import FPDFException as FPDFErr
except Exception:  # fallback if fpdf.errors isn't present
    class FPDFErr(Exception):
        pass
        
# ============================================================
# ================== NEW: VOICE IMPORTS ======================
# ============================================================
try:
    from streamlit_mic_recorder import mic_recorder
except ModuleNotFoundError:
    mic_recorder = None # Make it optional
try:
    import soundfile as sf
except ModuleNotFoundError:
    sf = None # Make it optional
try:
    import numpy as np
except ModuleNotFoundError:
    np = None # Make it optional
# ============================================================
# ============================================================
# ================ SENTIMENT ANALYSIS IMPORTS ================
# ============================================================
try:
    from transformers import pipeline
except ModuleNotFoundError:
    pipeline = None
try:
    import librosa
except ModuleNotFoundError:
    librosa = None
try:
    import pandas as pd
except ModuleNotFoundError:
    pd = None
# ============================================================
# ============================================================
# ================ VOCAL REMOVER IMPORTS ================
# ============================================================
try:
    from vocal_remover import nets, spec_utils
    from vocal_remover.separator import Separator
except ModuleNotFoundError:
    nets = None
    spec_utils = None
    Separator = None
# ============================================================

# Optional: hide Streamlit tracebacks in UI
st.set_option('client.showErrorDetails', False)


# ============================================================
# =============== API KEYS (as requested) ====================
# ============================================================



# ============================================================
# ================== DEVICE / COMPUTE TYPE ===================
# ============================================================
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

# ============================================================
# ================== GLOBAL VIDEO META =======================
# ============================================================
def empty_meta() -> Dict[str, Any]:
    return {
        "source": None,        # "youtube" | "file"
        "page_url": None,      # YouTube page url (for links)
        "file_path": None,     # local mp4 path (for VLC)
        "base": None,          # base name without extension
    }

# ============================================================
# ===================== UTILITIES ============================
# ============================================================
def sanitize_filename(name: str) -> str:
    name = re.sub(r'[<>:\\/|?*"]', "_", name)
    name = re.sub(r"[^0-9a-zA-Z _.-]", "_", name)
    name = re.sub(r"_+", "_", name)
    return name.strip(" ._")

def is_youtube_url(url: str) -> bool:
    pats = [
        r"(?:https?://)?(?:www\.)?youtube\.com/watch\?v=[\w-]+",
        r"(?:https?://)?(?:www\.)?youtu\.be/[\w-]+",
        r"(?:https?://)?(?:m\.)?youtube\.com/watch\?v=[\w-]+",
        r"(?:https?://)?(?:www\.)?youtube\.com/embed/[\w-]+",
    ]
    return any(re.match(p, url.strip()) for p in pats)

def _add_time_param(url: str, seconds: int) -> str:
    sep = "&" if "?" in url else "?"
    return f"{url}{sep}t={max(0,int(seconds))}s"

def _secs_to_hhmmss(s: Optional[int]) -> str:
    if s is None:
        return "??:??:??"
    return str(timedelta(seconds=int(s)))

# ============================================================
# ============ DOWNLOAD / EXTRACT AUDIO ======================
# ============================================================
def download_youtube_video(url: str) -> Optional[str]:
    ydl_opts = {
        "outtmpl": "%(title)s.%(ext)s",
        "merge_output_format": "mp4",
        "format": "bestvideo[ext=mp4]+bestaudio[ext=m4a]/best/best",
        "noplaylist": True,
        "quiet": True,
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            filename = ydl.prepare_filename(info)
            sanitized_title = sanitize_filename(info.get("title", "video"))
            final_path = f"{sanitized_title}.mp4"
            if os.path.exists(final_path):
                return os.path.abspath(final_path)
            elif os.path.exists(filename):
                return os.path.abspath(filename)
            else:
                return None
    except Exception as e:
        st.error(f"[yt_dlp] {e}")
        return None

def extract_audio(video_path: str) -> Optional[str]:
    base = os.path.splitext(video_path)[0]
    out = f"{base}_extracted_audio.wav"
    try:
        ffmpeg.input(video_path).output(out, acodec="pcm_s16le", ac=1, ar="16000").overwrite_output().run(quiet=True)
        return out
    except Exception as e:
        st.error(f"[ffmpeg] {e}")
        return None

def master_audio(audio_path: str, model_path='vocal-remover/models/baseline.pth') -> Optional[str]:
    """
    Separates the vocals from an audio file using the vocal-remover model.
    """
    st.info("Starting Vocal Separation to get Master Audio...")
    if nets is None or spec_utils is None or Separator is None:
        st.error("Vocal remover libraries not installed. Falling back to original audio.")
        return audio_path

    try:
        # Configuration
        sr = 44100
        n_fft = 2048
        hop_length = 1024
        batchsize = 4
        cropsize = 256
        
        st.info("Loading vocal separation model...")
        device = torch.device('cuda:0' if torch.cuda.is_available() else 'cpu')
        model = nets.CascadedNet(n_fft, hop_length, 32, 128)
        model.load_state_dict(torch.load(model_path, map_location=device))
        model.to(device)
        st.info("Model loaded successfully.")

        st.info(f"Loading audio file: {audio_path}")
        X, sr_loaded = librosa.load(
            audio_path, sr=sr, mono=False, dtype=np.float32, res_type='kaiser_fast'
        )
        if X.ndim == 1:
            X = np.asarray([X, X])

        st.info("Converting audio to spectrogram...")
        X_spec = spec_utils.wave_to_spectrogram(X, hop_length, n_fft)

        separator = Separator(model=model, device=device, batchsize=batchsize, cropsize=cropsize)
        
        st.info("Separating vocals from instruments...")
        _, v_spec = separator.separate_tta(X_spec)
        st.info("Separation complete.")

        st.info("Converting vocal spectrogram back to audio...")
        wave = spec_utils.spectrogram_to_wave(v_spec, hop_length=hop_length)

        base_name = os.path.splitext(os.path.basename(audio_path))[0]
        master_audio_path = f"{base_name}_master_vocals.wav"
        
        st.info(f"Saving master vocal track to: {master_audio_path}")
        sf.write(master_audio_path, wave.T, sr)
        
        st.success("Master audio (vocals) is ready.")
        return master_audio_path

    except Exception as e:
        st.error(f"An error occurred during vocal separation: {e}")
        return audio_path  # fallback

# ============================================================
# ============== OOM-SAFE WHISPERX LOADER ====================
# ============================================================
def _empty_cuda():
    try:
        torch.cuda.empty_cache()
    except Exception:
        pass
    gc.collect()

def _try_load_and_transcribe(model_name:str, device:str, compute_type:str, audio, batch_size:int):
    model = whisperx.load_model(model_name, device, compute_type=compute_type)
    res = model.transcribe(audio, batch_size=batch_size)
    del model
    _empty_cuda()
    return res

def transcribe_with_whisperx(audio_file: str) -> Optional[Dict]:
    st.write("**Transcribing (WhisperX) with OOM-safe ladder...**")
    audio = whisperx.load_audio(audio_file)

    attempts = []
    if DEVICE == "cuda":
        attempts += [
            ("large-v3",   "cuda", "int8_float16", 8),
            ("large-v2",   "cuda", "int8_float16", 8),
            ("medium.en",  "cuda", "int8_float16", 8),
            ("small.en",   "cuda", "int8_float16", 8),
            ("medium.en",  "cuda", "int8",         6),
            ("small.en",   "cuda", "int8",         6),
            ("small.en",   "cuda", "int8",         4),
        ]
    attempts += [
        ("medium.en",  "cpu",  "int8",  8),
        ("small.en",   "cpu",  "int8",  8),
    ]

    last_err = None
    for (mname, dev, ctype, bs) in attempts:
        try:
            with st.spinner(f"Trying model={mname} device={dev} compute_type={ctype} batch={bs}"):
                return _try_load_and_transcribe(mname, dev, ctype, audio, bs)
        except RuntimeError as e:
            last_err = e
            _empty_cuda()
        except Exception as e:
            last_err = e
            _empty_cuda()

    st.error("All WhisperX attempts failed.")
    if last_err:
        st.code(str(last_err))
    return None

def generate_chapters_with_llama(
    transcript: str, segments: List[Dict]
) -> Optional[List[Dict]]:
    """Generate chapters with timestamps using Llama 405B via Nebius"""
    nebius_client = OpenAI(base_url="https://api.studio.nebius.com/v1/", api_key=NEBIUS_API_KEY)
    if not nebius_client:
        st.error("Nebius client not available. Skipping chapter generation.")
        return None

    try:
        # Prepare transcript with timestamps for better chapter detection
        timestamped_transcript = ""
        for segment in segments:
            start_time = str(timedelta(seconds=int(segment["start"])))
            timestamped_transcript += f"[{start_time}] {segment['text']}\n"

        response = nebius_client.chat.completions.create(
            model="Qwen/Qwen3-235B-A22B-Instruct-2507",
            messages=[
                {
                    "role": "system",
                    "content": """You are an expert at analyzing video transcripts and creating meaningful chapters.
                        Create chapters that represent distinct topics or sections in the content.
                        Return the result as a JSON array with this format:
                        [{"title": "Chapter Title", "start_time": "00:00:00", "end_time": "00:05:30", "description": "Brief description"}]""",
                },
                {
                    "role": "user",
                    "content": f"Please analyze this timestamped transcript and create meaningful chapters:\n\n{timestamped_transcript}",
                },
            ],
            temperature=0.2,
            max_tokens=1500,
        )

        # Parse the JSON response
        chapters_text = response.choices[0].message.content
        if not chapters_text or not chapters_text.strip():
            st.error("Empty response from model. Cannot generate chapters.")
            return None
        import json
        json_start = chapters_text.find("[")
        json_end = chapters_text.rfind("]") + 1
        if json_start != -1 and json_end != -1:
            chapters_json = chapters_text[json_start:json_end]
            return json.loads(chapters_json)
        else:
            st.error("Could not parse chapters JSON from response")
            return None

    except Exception as e:
        st.error(f"Error generating chapters: {e}")
        return None

# ============================================================
# =================== WHISPERX ALIGN/DIAR ====================
# ============================================================
def get_diarization_pipeline(hf_token: str, device: str):
    if hasattr(whisperx, "load_diarize_model"):
        try:
            return whisperx.load_diarize_model(use_auth_token=hf_token, device=device)
        except Exception:
            pass
    try:
        import whisperx.diarize as wxd
        if hasattr(wxd, "DiarizationPipeline"):
            return wxd.DiarizationPipeline(use_auth_token=hf_token, device=device)
    except Exception:
        pass
    if hasattr(whisperx, "DiarizationPipeline"):
        try:
            return whisperx.DiarizationPipeline(use_auth_token=hf_token, device=device)
        except Exception:
            pass
    return None

def diarize_with_whisperx(audio_file: str, transcription_result: Dict) -> Optional[Dict]:
    st.write("**Diarizing (WhisperX)...**")
    lang = transcription_result.get("language")
    model_a, metadata = whisperx.load_align_model(language_code=lang, device=DEVICE)
    result_aligned = whisperx.align(transcription_result["segments"], model_a, metadata, audio_file, DEVICE)
    del model_a; gc.collect(); _empty_cuda()

    diarize_model = get_diarization_pipeline(HF_TOKEN, DEVICE)
    if diarize_model is None:
        raise RuntimeError("No diarization API found in your WhisperX build.")

    diarize_segments = diarize_model(audio_file)
    del diarize_model; gc.collect(); _empty_cuda()

    return whisperx.assign_word_speakers(diarize_segments, result_aligned)

def format_diarized_transcript(result: Dict, output_path: str):
    def fmt(seconds: float) -> str:
        if seconds is None: return "??:??:??"
        return str(timedelta(seconds=int(seconds)))
    lines = []
    if "segments" in result:
        for seg in result["segments"]:
            start = fmt(seg.get("start"))
            end = fmt(seg.get("end"))
            speaker = seg.get("speaker", "SPEAKER_UNKNOWN")
            text = (seg.get("text") or "").strip()
            lines.append(f"[{start} - {end}] {speaker}: {text}")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
# ============================================================
# ============= SENTIMENT ANALYSIS (INTEGRATED) ==============
# ============================================================
loaded_pipelines = {}

@st.cache_resource
def get_sota_pipeline(task, model_name):
    """Loads and caches a Hugging Face pipeline to avoid reloading."""
    if pipeline is None:
        st.error("`transformers` library not found. Please install it: `pip install transformers`")
        return None
        
    # This check is now redundant due to @st.cache_resource, but safe to keep
    if model_name in loaded_pipelines:
        return loaded_pipelines[model_name]
    else:
        progress_text = f"Loading SOTA expert model for {task}: {model_name}..."
        st.info(progress_text)
        
        # Use Streamlit's caching for the pipeline object itself
        pipeline_instance = pipeline(
            task, model=model_name, return_all_scores=True,
            device=0 if DEVICE == "cuda" else -1
        )
        loaded_pipelines[model_name] = pipeline_instance
        st.info(f"✅ {model_name} ready.")
        return pipeline_instance

def run_sota_emotion_analysis(audio_path: str, whisper_result: dict, start_time: float, end_time: float, analysis_type: str) -> Optional[Dict]:
    """
    Performs emotion analysis and returns a dictionary of results for UI display.
    """
    st.info(f"Analyzing segment from {start_time:.2f}s to {end_time:.2f}s...")

    try:
        # --- 1. Extract Text and Audio Segments ---
        text_segment = ""
        for seg in whisper_result.get("segments", []):
            if seg['start'] < end_time and seg['end'] > start_time:
                text_segment += seg['text'] + " "
        text_segment = text_segment.strip()

        if not text_segment:
            st.warning("Could not find any transcribed text in the selected time range.")
            # We can still proceed if it's a speech-only analysis

        st.info(f"Text in segment: \"{text_segment}\"")
        
        # Load audio segment if required
        audio_segment = None
        if analysis_type in ['Speech-Only', 'Multimodal']:
            if librosa is None:
                st.error("`librosa` library not found. Please install it for audio analysis: `pip install librosa`")
                return None
            st.info("Loading audio segment...")
            y, sr = librosa.load(audio_path, sr=16000)
            start_sample = int(start_time * sr)
            end_sample = int(end_time * sr)
            audio_segment = y[start_sample:end_sample]
            st.info(f"Audio segment loaded: {len(audio_segment)} samples")

        # --- 2. Define Models and Label Mappings ---
        text_expert_model = "cardiffnlp/twitter-roberta-base-emotion"
        speech_expert_model = "superb/wav2vec2-base-superb-er"
        
        COMMON_LABELS = ['Positive', 'Negative', 'Neutral', 'Mixed/Other']
        text_label_map = {'joy': 'Positive', 'optimism': 'Positive', 'sadness': 'Negative', 'anger': 'Negative'}
        speech_label_map = {'hap': 'Positive', 'sad': 'Negative', 'ang': 'Negative', 'neu': 'Neutral'}

        def process_scores(raw_scores, label_map):
            processed = {label: 0.0 for label in COMMON_LABELS}
            for score_dict in raw_scores:
                model_label = score_dict['label']
                if model_label in label_map:
                    common_label = label_map[model_label]
                    processed[common_label] += score_dict['score']
            return processed

        # --- 3. Execute Analysis Based on User's Choice ---
        final_scores = {label: 0.0 for label in COMMON_LABELS}
        
        if analysis_type == 'Text-Only':
            if not text_segment:
                st.error("No text found in segment for Text-Only analysis.")
                return None
            text_classifier = get_sota_pipeline("text-classification", text_expert_model)
            text_scores_raw = text_classifier(text_segment)[0]
            final_scores = process_scores(text_scores_raw, text_label_map)

        elif analysis_type == 'Speech-Only':
            if audio_segment is None or len(audio_segment) == 0:
                st.error("Audio segment is empty for Speech-Only analysis.")
                return None
            speech_recognizer = get_sota_pipeline("audio-classification", speech_expert_model)
            speech_scores_raw = speech_recognizer(audio_segment)
            final_scores = process_scores(speech_scores_raw, speech_label_map)

        elif analysis_type == 'Multimodal':
            if not text_segment:
                st.warning("No text found for multimodal; analysis will be speech-only.")
            if audio_segment is None or len(audio_segment) == 0:
                st.error("Audio segment is empty for Multimodal analysis.")
                return None

            # Text analysis
            text_scores = {l: 0.0 for l in COMMON_LABELS}
            if text_segment:
                text_classifier = get_sota_pipeline("text-classification", text_expert_model)
                text_scores_raw = text_classifier(text_segment)[0]
                text_scores = process_scores(text_scores_raw, text_label_map)
            
            # Speech analysis
            speech_recognizer = get_sota_pipeline("audio-classification", speech_expert_model)
            speech_scores_raw = speech_recognizer(audio_segment)
            speech_scores = process_scores(speech_scores_raw, speech_label_map)

            # Late Fusion with Weighted Average
            text_weight = 0.5
            speech_weight = 0.5
            for label in COMMON_LABELS:
                final_scores[label] = (text_scores.get(label, 0.0) * text_weight) + \
                                      (speech_scores.get(label, 0.0) * speech_weight)
        
        # --- 4. Prepare Results for Display ---
        total_score = sum(final_scores.values())
        if total_score < 1e-6:
            predicted_emotion = "Mixed/Other"
            note = "The detected emotions did not map to Positive/Negative/Neutral."
        else:
            predicted_emotion = max(final_scores, key=final_scores.get)
            note = None
        
        return {
            "dominant": predicted_emotion,
            "scores": final_scores,
            "total_score": total_score,
            "note": note
        }

    except Exception as e:
        st.error(f"An unexpected error occurred during emotion analysis: {e}")
        return None

# ============================================================
# =============== ADVANCED SUMMARIZATION (MAP-REDUCE) ========
# ============================================================
class SummarizationGoal(Enum):
    MEETING_MINUTES = "meeting_minutes"
    PODCAST_SUMMARY = "podcast_summary"
    LECTURE_NOTES = "lecture_notes"
    INTERVIEW_HIGHLIGHTS = "interview_highlights"
    GENERAL_SUMMARY = "general_summary"

def _nebius_client():
    return OpenAI(base_url="https://api.studio.nebius.com/v1/", api_key=NEBIUS_API_KEY)

def split_text_into_chunks(text: str, max_words: int = 1500, overlap_sentences: int = 2) -> list:
    paragraphs = text.split("\n\n")
    all_sentences = []
    for p in paragraphs:
        sentences = re.split(r"(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s", p)
        all_sentences.extend(s for s in sentences if s)
    if not all_sentences:
        return []
    chunks, current_chunk_sentences, current_word_count = [], [], 0
    for sentence in all_sentences:
        sentence_word_count = len(sentence.split())
        if current_word_count + sentence_word_count > max_words and current_chunk_sentences:
            chunks.append(" ".join(current_chunk_sentences))
            overlap_index = max(0, len(current_chunk_sentences) - overlap_sentences)
            current_chunk_sentences = current_chunk_sentences[overlap_index:]
            current_word_count = sum(len(s.split()) for s in current_chunk_sentences)
        current_chunk_sentences.append(sentence)
        current_word_count += sentence_word_count
    if current_chunk_sentences:
        chunks.append(" ".join(current_chunk_sentences))
    return chunks

def call_nebius_llm(prompt: str,
                     model_name: str = "Qwen/Qwen3-235B-A22B-Instruct-2507",
                     system_prompt: str = "You are a helpful assistant.") -> Optional[str]:
    try:
        client = _nebius_client()
        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt},
            ],
            max_tokens=2048,
            temperature=0.2,
            top_p=0.9,
            extra_body={"top_k": 50},
        )
        return response.choices[0].message.content
    except Exception:
        return None

def extract_topics_and_keywords(summary_content: str, goal: SummarizationGoal) -> Optional[List[Dict]]:
    prompt = (
        f"Based on the following {goal.value.replace('_',' ') }:\n---\n{summary_content}\n---\n"
        f"Extract main topics. For each, provide a short title and 3-5 keywords. "
        f"Return ONLY a valid JSON array of objects."
    )
    response_text = call_nebius_llm(
        prompt,
        system_prompt="You are an expert at topic extraction and respond only with valid JSON."
    )
    if not response_text:
        return None
    try:
        import json
        json_match = re.search(r"\[.*\]", response_text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group())
        return None
    except Exception:
        return None

def summarize_long_transcript_map_reduce(transcript: str,
                                         goal: SummarizationGoal,
                                         summary_length_words: int = 250) -> Optional[Dict]:
    chunks = split_text_into_chunks(transcript)
    if not chunks:
        return None

    # Map
    map_prompts = {
        SummarizationGoal.MEETING_MINUTES: "Summarize key decisions, action items, and main points in THIS meeting segment.",
        SummarizationGoal.PODCAST_SUMMARY: "Summarize the main topics and key takeaways from THIS part of the podcast.",
        SummarizationGoal.LECTURE_NOTES: "Create bullet-point notes of key concepts and definitions from THIS lecture section.",
        SummarizationGoal.INTERVIEW_HIGHLIGHTS: "Extract insightful points and notable quotes from THIS part of the interview.",
        SummarizationGoal.GENERAL_SUMMARY: "Provide a concise summary of main points in this transcript segment.",
    }
    chunk_summaries = []
    for i, chunk in enumerate(chunks):
        summary = call_nebius_llm(
            prompt=f"{map_prompts[goal]}\n\nTranscript Segment:\n---\n{chunk}",
            system_prompt="You are an expert at summarizing professional content.",
        )
        if summary:
            chunk_summaries.append(summary)

    if not chunk_summaries:
        return None

    # Reduce with length control
    transcript_word_count = len(transcript.split())
    upper_bound_words = max(50, int(transcript_word_count * 0.40))
    if summary_length_words > upper_bound_words:
        summary_length_words = upper_bound_words

    combined_summaries = "\n\n---\n\n".join(chunk_summaries)
    reduce_prompts = {
        SummarizationGoal.MEETING_MINUTES: "Synthesize into a report with sections for 'Overall Summary', 'Main Topics', 'Key Decisions', and 'Action Items'.",
        SummarizationGoal.PODCAST_SUMMARY: "Create show notes with 'Episode Title', 'Introduction', 'Key Segments', and 'Memorable Quotes'.",
        SummarizationGoal.LECTURE_NOTES: "Organize into lecture notes with nested bullet points for 'Main Concepts'.",
        SummarizationGoal.INTERVIEW_HIGHLIGHTS: "Create a highlights report with 'Synopsis', 'Key Insights', and 'Notable Quotes'.",
        SummarizationGoal.GENERAL_SUMMARY: "Create a report with 'Overall Summary', 'Main Topics', and 'Key Points'.",
    }
    length_instruction = f"\n\nIMPORTANT FINAL INSTRUCTION: The entire final report MUST be approximately {summary_length_words} words long. Be concise."
    final_prompt = f"{reduce_prompts[goal]}\n\nHere are the summaries:\n---\n{combined_summaries}\n---{length_instruction}"

    final_summary = call_nebius_llm(
        prompt=final_prompt,
        system_prompt="You are an expert analyst creating a final report.",
    )
    if not final_summary:
        return None

    topics = extract_topics_and_keywords(final_summary, goal)
    return {"final_report": final_summary, "topics": topics}

# ============================================================
# ============== YT Official Transcript Fetcher ==============
# ============================================================
def fetch_official_youtube_transcript(youtube_url: str) -> Tuple[bool, Optional[str], Optional[List[Dict[str, Any]]]]:
    # returns (has_official, text, raw_items)
    def get_video_id(url: str) -> Optional[str]:
        if "v=" in url:
            return url.split("v=")[1].split("&")[0]
        if "youtu.be/" in url:
            return url.split("youtu.be/")[1].split("?")[0]
        if "embed/" in url:
            return url.split("embed/")[1].split("?")[0]
        return None

    vid = get_video_id(youtube_url)
    if not vid:
        return False, None, None
    try:
        tlist = YouTubeTranscriptApi.list_transcripts(vid)
        # prefer manually-created en, else generated en
        try:
            tr = tlist.find_manually_created_transcript(['en'])
        except Exception:
            tr = tlist.find_generated_transcript(['en'])
        data = tr.fetch()
        text = " ".join([t["text"] for t in data if t["text"].strip()])
        return True, text, data
    except (NoTranscriptFound, TranscriptsDisabled):
        return False, None, None
    except Exception:
        return False, None, None

# ============================================================
# ================= SENTENCE-LEVEL RAG =======================
# ============================================================
_TS_RE = re.compile(r"^\[(\d+):(\d{2}):(\d{2})\s*-\s*(\d+):(\d{2}):(\d{2})\]\s*(.*?):\s*(.*)$")

def _hms_to_seconds(h, m, s): return int(h)*3600 + int(m)*60 + int(s)

def parse_diarized_lines(path: str) -> List[Dict[str, Any]]:
    items = []
    with open(path, "r", encoding="utf-8") as f:
        for ln in f:
            ln = ln.strip()
            if not ln: continue
            m = _TS_RE.match(ln)
            if not m:
                items.append({"text": ln, "speaker": "UNKNOWN", "start": None, "end": None, "stamp": None})
                continue
            h1, m1, s1, h2, m2, s2, speaker, text = m.groups()
            start = _hms_to_seconds(h1, m1, s1)
            end = _hms_to_seconds(h2, m2, s2)
            items.append({
                "text": text.strip(),
                "speaker": speaker.strip(),
                "start": start, "end": end,
                "stamp": f"[{h1}:{m1}:{s1} - {h2}:{m2}:{s2}]"
            })
    return items

class SentenceRAG:
    def __init__(self, vectorstore: FAISS, embedder: HuggingFaceEmbeddings, entries: List[Dict[str, Any]], video_meta: Dict[str, Any]):
        self.vs = vectorstore
        self.embedder = embedder
        self.entries = entries
        self.video_meta = video_meta

        self.llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0, max_tokens=220, request_timeout=25)
        self.retriever = self.vs.as_retriever(search_type="mmr", search_kwargs={"k": 9, "fetch_k": 36, "lambda_mult": 0.25})
        self.prompt = PromptTemplate.from_template(
            "Answer ONLY from the context.\n"
            "Include the most relevant timestamp at the end as (Source: [HH:MM:SS - HH:MM:SS]).\n"
            "If not in context, reply exactly: I cannot answer the question based on the provided context.\n\n"
            "CONTEXT:\n{context}\n\nQUESTION: {q}\n\nANSWER:"
        )
        self.chain = create_stuff_documents_chain(self.llm, self.prompt)

    @staticmethod
    def _start_key(d: Document):
        st = d.metadata.get("start")
        return 10**12 if st is None else st

    def _compact_context(self, docs: List[Document], max_chars: int = 1200) -> str:
        docs_sorted = sorted(docs, key=self._start_key)
        seen = set()
        out, total = [], 0
        for d in docs_sorted:
            key = (d.page_content.strip(), d.metadata.get("start"), d.metadata.get("end"))
            if key in seen: continue
            seen.add(key)
            stamp = d.metadata.get("stamp") or ""
            line = f"{stamp} {d.page_content.strip()}"
            if total + len(line) > max_chars: break
            out.append(line); total += len(line) + 1
        return "\n".join(out)

    def _build_segment_link(self, start_seconds: Optional[int]) -> Optional[str]:
        if start_seconds is None:
            return None
        if self.video_meta.get("source") == "youtube" and self.video_meta.get("page_url"):
            return _add_time_param(self.video_meta["page_url"], start_seconds)
        if self.video_meta.get("source") == "file" and self.video_meta.get("file_path"):
            return f'vlc --start-time={int(start_seconds)} "{self.video_meta["file_path"]}"'
        return None

    def answer(self, question: str) -> Tuple[str, Optional[str]]:
        docs_mmr = self.retriever.get_relevant_documents(question)
        if not docs_mmr:
            return "I cannot answer the question based on the provided context.", None

        best_doc = docs_mmr[0]
        context = self._compact_context(docs_mmr)
        doc_for_llm = Document(page_content=context, metadata={})
        out = self.chain.invoke({"context": [doc_for_llm], "q": question})
        ans = (out.get("answer") if isinstance(out, dict) else out).strip()

        start = best_doc.metadata.get("start")
        end = best_doc.metadata.get("end")
        stamp = f"[{_secs_to_hhmmss(start)} - {_secs_to_hhmmss(end)}]"
        link = self._build_segment_link(start)
        tail = f"(Source: {stamp})"
        if link:
            tail += f" (Link: {link})"

        if "Source:" not in ans:
            ans = ans.rstrip() + " " + tail
        else:
            if link and "(Link:" not in ans:
                ans = ans.rstrip() + f" (Link: {link})"
        return ans, link

def initialize_sentence_rag(transcript_path: str, index_dir: str, video_meta: Dict[str, Any]) -> Optional[SentenceRAG]:
    try:
        if os.path.isdir(index_dir):
            embedder = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2", encode_kwargs={"normalize_embeddings": True})
            vs = FAISS.load_local(index_dir, embedder, allow_dangerous_deserialization=True)
            entries = parse_diarized_lines(transcript_path)
            return SentenceRAG(vs, embedder, entries, video_meta)

        entries = parse_diarized_lines(transcript_path)
        if not entries:
            st.error("Empty transcript."); return None

        docs = []
        for e in entries:
            if not e["text"]: continue
            docs.append(Document(page_content=e["text"],
                                 metadata={"speaker": e["speaker"], "start": e["start"], "end": e["end"], "stamp": e["stamp"]}))

        embedder = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2", encode_kwargs={"normalize_embeddings": True})
        vs = FAISS.from_documents(docs, embedder)

        os.makedirs(index_dir, exist_ok=True)
        vs.save_local(index_dir)
        return SentenceRAG(vs, embedder, entries, video_meta)
    except Exception as e:
        st.error(f"[RAG init] {e}")
        return None
# ============================================================
# ================== REAL-TIME VOICE Q&A =====================
# ============================================================

@st.cache_resource
def load_fast_transcriber():
    """
    Loads a WhisperX model into cache for voice queries.
    We try 'small.en' for better accuracy, falling back to 'base.en' if needed.
    """
    try:
        # 'small.en' is more accurate and a good balance for this task
        model = whisperx.load_model("small.en", DEVICE, compute_type="int8", language="en")
        st.success("Loaded 'small.en' model for voice queries.")
        return model
    except Exception as e:
        st.warning(f"Could not load 'small.en' model ({e}), falling back to 'base.en'.")
        try:
            # Fallback to the faster but less accurate model
            model = whisperx.load_model("base.en", DEVICE, compute_type="int8", language="en")
            st.success("Loaded 'base.en' model for voice queries.")
            return model
        except Exception as e2:
            st.error(f"Failed to load any voice model: {e2}")
            return None

def transcribe_voice_query(audio_info: dict) -> Optional[str]:
    """
    Transcribes audio bytes from the mic recorder.
    """
    if not audio_info or not audio_info.get("bytes"):
        st.warning("No audio data received from microphone")
        return None

    audio_bytes = audio_info["bytes"]
    sample_rate = audio_info.get("sample_rate", 44100)
    
    # Debug info
    st.sidebar.write(f"Audio info: {len(audio_bytes)} bytes at {sample_rate}Hz")
    
    # Save audio for debugging and processing
    debug_audio_path = "debug_last_voice_query.wav"
    
    try:
        # Convert bytes to numpy array
        if isinstance(audio_bytes, bytes):
            # Assume 16-bit PCM audio
            audio_np = np.frombuffer(audio_bytes, dtype=np.int16)
        else:
            audio_np = np.array(audio_bytes, dtype=np.int16)
        
        # Check if audio is not empty/silent
        if len(audio_np) == 0:
            st.warning("Empty audio recording received")
            return None
            
        # Check for silence (very low amplitude)
        max_amplitude = np.max(np.abs(audio_np))
        if max_amplitude < 100:  # Threshold for detecting silence
            st.warning("Audio appears to be silent or very quiet. Try speaking louder.")
            return None
            
        # Normalize audio if it's too quiet
        if max_amplitude > 0:
            audio_np = audio_np.astype(np.float32)
            audio_np = audio_np / max_amplitude * 0.7  # Normalize to 70% max
            audio_np = (audio_np * 32767).astype(np.int16)
        
        # Save to WAV file
        sf.write(debug_audio_path, audio_np, samplerate=sample_rate)
        
        # Display the audio player in the sidebar for immediate feedback
        st.sidebar.subheader("Last Voice Query Audio:")
        st.sidebar.audio(debug_audio_path)
        st.sidebar.write(f"Audio length: {len(audio_np)/sample_rate:.2f} seconds")
        
    except Exception as e:
        st.error(f"Error processing audio data: {e}")
        return None
    
    try:
        model = load_fast_transcriber()
        if model is None:
            return "Voice model not loaded. Please check the setup."
            
        # Load audio using WhisperX
        audio = whisperx.load_audio(debug_audio_path)
        
        # Check audio length
        if len(audio) < sample_rate * 0.5:  # Less than 0.5 seconds
            return "Audio too short. Please record for at least 1 second."
        
        # Transcribe with error handling
        result = model.transcribe(audio, batch_size=4)

        if result and result.get("segments"):
            segments_text = []
            for seg in result["segments"]:
                text = seg.get("text", "").strip()
                if text:
                    segments_text.append(text)
            
            full_text = " ".join(segments_text).strip()
            
            if full_text:
                st.sidebar.success(f"Transcribed: '{full_text}'")
                return full_text
            else:
                return "No speech detected in the audio. Please try speaking more clearly."
        else:
            return "Failed to transcribe audio. Please try again."
    
    except Exception as e:
        st.error(f"Error during voice transcription: {e}")
        return "Transcription failed. Please try again."

# ============================================================
# ===================== FLASHCARD GEN ========================
# ============================================================
def generate_flashcards_from_text(base_text: str, n_cards: int = 20) -> List[Dict[str, str]]:
    """
    Use Nebius LLM to generate concise Q/A flashcards.
    Returns list of dicts: [{"question": "...", "answer": "..."}]
    """
    prompt = (
        "Create high-quality study flashcards (Q/A pairs) from the content below.\n"
        f"Return EXACTLY {n_cards} JSON objects in an array, each with keys 'question' and 'answer'. "
        "Keep questions crisp and answers concise (1-3 sentences). "
        "Avoid duplicates, avoid filler. Respond ONLY with valid JSON.\n\n"
        "CONTENT:\n"
        f"{base_text[:18000]}\n"
    )
    txt = call_nebius_llm(prompt, system_prompt="You are a meticulous flashcard generator that outputs only JSON.")
    if not txt:
        return []
    import json
    m = re.search(r"\[.*\]", txt, re.DOTALL)
    if not m:
        return []
    try:
        arr = json.loads(m.group())
        cleaned = []
        for item in arr:
            q = (item.get("question") or "").strip()
            a = (item.get("answer") or "").strip()
            if q and a:
                cleaned.append({"question": q, "answer": a})
        return cleaned
    except Exception:
        return []

def flashcards_to_csv(fcards: List[Dict[str, str]]) -> bytes:
    """
    CSV writer expects text; we use StringIO then encode to bytes.
    """
    import csv, io
    s = io.StringIO(newline="")
    writer = csv.writer(s)
    writer.writerow(["question", "answer"])
    for c in fcards:
        writer.writerow([c["question"], c["answer"]])
    # UTF-8 with BOM helps Excel open properly
    return ("\ufeff" + s.getvalue()).encode("utf-8")

def flashcards_to_tsv(fcards: List[Dict[str, str]]) -> bytes:
    out = []
    for c in fcards:
        out.append(f"{c['question']}\t{c['answer']}")
    return ("\n".join(out) + "\n").encode("utf-8")

# ============================================================
# ============== PDF SAFE WRAPPING HELPERS ===================
# ============================================================
def _break_long_words(text: str, max_word_len: int = 60) -> str:
    """
    Prevent FPDF 'Not enough horizontal space' by inserting hard breaks
    inside long unbreakable tokens (e.g., giant URLs/hashes).
    """
    def chunk_token(tok: str) -> str:
        if len(tok) <= max_word_len:
            return tok
        return "\n".join(tok[i:i+max_word_len] for i in range(0, len(tok), max_word_len))

    out_lines = []
    for line in text.splitlines():
        # Replace non-breaking spaces & zero-width chars with normal spaces
        line = (line
                .replace("\u00A0", " ")
                .replace("\u202F", " ")
                .replace("\u2007", " ")
                .replace("\u2009", " ")
                .replace("\u200A", " ")
                .replace("\u200B", "")  # zero-width space
                .replace("\u2060", " ")
                .replace("\t", "    "))
        # preserve whitespace separators
        parts = re.split(r"(\s+)", line)
        parts = [chunk_token(p) if not p.isspace() else p for p in parts]
        out_lines.append("".join(parts))
    return "\n".join(out_lines)

def _sanitize_for_pdf_ascii(text: str) -> str:
    """
    Fallback for core fonts: strip/replace to Latin-1 friendly set.
    """
    text = text.replace("\t", "    ")
    text = _break_long_words(text, max_word_len=70)
    # collapse absurd runs of symbols
    text = re.sub(r"([-_])\1{79,}", r"\1" * 79, text)
    # force to latin-1 (replace unknowns)
    return text.encode("latin-1", "replace").decode("latin-1")

def _find_ttf_font() -> Optional[str]:
    """
    Try to locate a Unicode TTF on common Linux/macOS paths.
    """
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        os.path.join(os.path.dirname(__file__) if '__file__' in globals() else ".", "DejaVuSans.ttf"),
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None

def safe_multicell(pdf: FPDF, width: float, height: float, text: str):
    """
    Multicell wrapper that never raises into Streamlit.
    Ensures left margin X, and falls back to chunked/ASCII on errors.
    """
    # hard-break long tokens and normalize odd spaces
    txt = _break_long_words(text, max_word_len=80)
    # ensure we're starting at left margin every time
    pdf.set_x(pdf.l_margin)
    try:
        pdf.multi_cell(width, height, txt)
        return
    except FPDFErr:
        pass

    # Fallback: try again with latin-1 replacement
    try:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(width, height, txt.encode("latin-1", "replace").decode("latin-1"))
        return
    except FPDFErr:
        pass

    # Last resort: hard chunk by 80 chars so a single "character" never exceeds width
    for chunk in [txt[i:i+80] for i in range(0, len(txt), 80)]:
        try:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(width, height, chunk)
        except Exception:
            # absolute last fallback: one line cell with latin-1
            pdf.set_x(pdf.l_margin)
            pdf.cell(width, height, chunk.encode("latin-1", "replace").decode("latin-1"))
            pdf.ln(height)

# ============================================================
# ===================== CONVO EXPORTS ========================
# ============================================================
def conversation_to_markdown(chat_history: List[Tuple[str, str]]) -> str:
    md = "# Conversation Export\n\n"
    for role, msg in chat_history:
        name = "You" if role == "user" else "Assistant"
        md += f"**{name}:**\n\n{msg}\n\n---\n\n"
    return md

def conversation_to_docx_bytes(chat_history: List[Tuple[str, str]]) -> Optional[bytes]:
    if DOCXDocument is None:
        return None
    doc = DOCXDocument()
    doc.add_heading("Conversation Export", level=1)
    for role, msg in chat_history:
        name = "You" if role == "user" else "Assistant"
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        doc.add_paragraph(msg)
        doc.add_paragraph("---")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def conversation_to_pdf_bytes(chat_history: List[Tuple[str, str]]) -> Optional[bytes]:
    if FPDF is None:
        return None
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Prefer a Unicode TTF if available (robust for emojis/intl chars)
    ttf = _find_ttf_font()
    if ttf:
        pdf.add_font("UF", "", ttf, uni=True)
        header_font = ("UF", "", 13)
        label_font  = ("UF", "", 12)
        body_font   = ("UF", "", 11)
    else:
        header_font = ("Arial", "", 13)
        label_font  = ("Arial", "", 12)
        body_font   = ("Arial", "", 11)

    usable_w = pdf.w - pdf.l_margin - pdf.r_margin

    pdf.set_font(*header_font)
    safe_multicell(pdf, usable_w, 8, "Conversation Export")
    pdf.ln(2)

    for role, msg in chat_history:
        name = "You" if role == "user" else "Assistant"

        pdf.set_font(*label_font)
        safe_multicell(pdf, usable_w, 6, f"{name}:")   # no width=0 anywhere
        pdf.set_font(*body_font)

        # choose sanitizer based on font availability
        if ttf:
            safe_msg = (msg.replace("\u00A0", " ").replace("\u202F", " ").replace("\u2007", " ")
                           .replace("\u2009", " ").replace("\u200A", " ").replace("\u200B", "")
                           .replace("\u2060", " ").replace("\t", "    "))
            safe_msg = _break_long_words(safe_msg, max_word_len=80)
        else:
            safe_msg = _sanitize_for_pdf_ascii(msg)

        # render line-by-line with safe_multicell
        for line in (safe_msg.splitlines() or [""]):
            if not line.strip():
                pdf.ln(4)
            else:
                safe_multicell(pdf, usable_w, 5.5, line)

        pdf.ln(1)
        safe_multicell(pdf, usable_w, 5, "-" * 40)

    # Return bytes safely (avoid raising into Streamlit)
    try:
        return pdf.output(dest="S").encode("latin-1", "ignore")
    except Exception:
        return None

# ============================================================
# ======================== STREAMLIT UI ======================
# ============================================================
st.set_page_config(page_title="Fast RAG Q&A (WhisperX) + Map-Reduce Summarization", layout="wide")
st.title("🎧 Fast RAG Q&A with WhisperX, Diarization & Timestamped Deep Links")
st.caption("Includes official YouTube transcript (if available), Nebius map-reduce summarization, exports, and flashcards.")

# Session state
if "video_meta" not in st.session_state:
    st.session_state.video_meta = empty_meta()
if "diar_path" not in st.session_state:
    st.session_state.diar_path = None
if "index_dir" not in st.session_state:
    st.session_state.index_dir = None
if "qna" not in st.session_state:
    st.session_state.qna = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "yt_transcript_available" not in st.session_state:
    st.session_state.yt_transcript_available = False
if "yt_transcript_text" not in st.session_state:
    st.session_state.yt_transcript_text = None

with st.sidebar:
    st.header("1) Input")
    src_choice = st.radio("Source", ["YouTube URL", "Upload local file"], horizontal=False)

    if src_choice == "YouTube URL":
        yt_url = st.text_input("YouTube URL", placeholder="https://www.youtube.com/watch?v=...")
        load_btn = st.button("Download / Load")
        # Show a helper to fetch official transcript on demand
        fetch_tx_btn = st.button("Fetch Official YT Transcript")
    else:
        uploaded = st.file_uploader("Upload video file (mp4/mov/mkv)", type=["mp4","mov","mkv"])
        load_btn = st.button("Save Upload")
        fetch_tx_btn = False  # not applicable

    st.divider()
    st.header("2) Build Data")
    master_audio_enabled = st.checkbox("Master Audio (enhance quality)", value=False)
    tnd_btn = st.button("Transcribe + Diarize (one-time)")
    generate_chapters_btn = st.button("Generate Chapters (optional)")

    st.divider()
    st.header("3) Summarization (Nebius map-reduce)")
    use_official_yt = st.checkbox("Prefer official YouTube transcript when available", value=True)
    goal_map = {
        "General Summary": SummarizationGoal.GENERAL_SUMMARY,
        "Meeting Minutes": SummarizationGoal.MEETING_MINUTES,
        "Podcast Summary": SummarizationGoal.PODCAST_SUMMARY,
        "Lecture Notes": SummarizationGoal.LECTURE_NOTES,
        "Interview Highlights": SummarizationGoal.INTERVIEW_HIGHLIGHTS,
    }
    sum_goal_label = st.selectbox("Summary goal", list(goal_map.keys()))
    sum_len = st.slider("Target summary length (words)", 100, 1000, 300, 50)
    run_summary_btn = st.button("Run Summarization Now")

    st.divider()
    st.header("4) RAG Index")
    build_btn = st.button("Build / Load RAG Index")
 

    st.divider()
    st.header("Utilities")
    reset_btn = st.button("Reset Session")

if reset_btn:
    # Clear all session state
    keys_to_clear = list(st.session_state.keys())
    for key in keys_to_clear:
        del st.session_state[key]
    
    # Reinitialize essential state
    st.session_state.video_meta = empty_meta()
    st.session_state.chat_history = []
    st.session_state.yt_transcript_available = False
    st.session_state.yt_transcript_text = None
    st.rerun()

# 1) Load video
if 'load_btn' in locals() and load_btn:
    with st.spinner("Preparing video..."):
        if src_choice == "YouTube URL":
            if not yt_url or not is_youtube_url(yt_url):
                st.error("Enter a valid YouTube URL.")
            else:
                vpath = download_youtube_video(yt_url)
                if not vpath:
                    st.error("Failed to download video.")
                else:
                    st.success("Video downloaded.")
                    st.session_state.video_meta = empty_meta()
                    st.session_state.video_meta["source"] = "youtube"
                    st.session_state.video_meta["file_path"] = vpath
                    st.session_state.video_meta["base"] = os.path.splitext(vpath)[0]
                    st.session_state.video_meta["page_url"] = yt_url
        else:
            if not uploaded:
                st.error("Please upload a file.")
            else:
                suffix = os.path.splitext(uploaded.name)[1]
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(uploaded.read())
                    vpath = tmp.name
                st.success(f"Saved to {vpath}")
                st.session_state.video_meta = empty_meta()
                st.session_state.video_meta["source"] = "file"
                st.session_state.video_meta["file_path"] = vpath
                st.session_state.video_meta["base"] = os.path.splitext(vpath)[0]

# Fetch official YT transcript (on-demand button)
if fetch_tx_btn:
    vm = st.session_state.video_meta
    if vm.get("source") == "youtube" and (vm.get("page_url") or 'yt_url' in locals()):
        page = vm.get("page_url") or yt_url
        with st.spinner("Fetching official YouTube transcript..."):
            ok, txt, raw = fetch_official_youtube_transcript(page)
        if ok and txt:
            st.session_state.yt_transcript_available = True
            st.session_state.yt_transcript_text = txt
            st.success(f"Fetched official transcript. Length: {len(txt.split())} words")
            st.text_area("Transcript preview (first 2,000 chars)", txt[:2000], height=200)
            st.download_button("Download transcript.txt", data=txt.encode("utf-8"), file_name="transcript.txt")
            # also generate .srt from raw
            if raw:
                def srt_time(s):
                    s = float(s)
                    h = int(s // 3600); m = int((s % 3600) // 60); sec = int(s % 60)
                    ms = int((s - int(s)) * 1000)
                    return f"{h:02}:{m:02}:{sec:02},{ms:03}"
                lines = []
                for idx, i in enumerate(raw, 1):
                    start = i["start"]
                    end = i["start"] + i.get("duration", 0)
                    text = i["text"].replace("\n", " ").strip()
                    lines += [str(idx), f"{srt_time(start)} --> {srt_time(end)}", text, ""]
                srt_bytes = "\n".join(lines).encode("utf-8")
                st.download_button("Download transcript.srt", data=srt_bytes, file_name="transcript.srt")
        else:
            st.warning("Official transcript unavailable for this video.")
    else:
        st.info("Load a YouTube video first.")

# 2) Transcribe + Diarize
vm = st.session_state.video_meta
if vm.get("file_path"):
    st.info(f"Loaded: `{vm['file_path']}`")
    if vm.get("source") == "youtube" and vm.get("page_url"):
        st.write(f"Page: {vm['page_url']}")

if tnd_btn:
    if not vm.get("file_path"):
        st.error("Please load a video first.")
    else:
        with st.spinner("Extracting audio..."):
            apath = extract_audio(vm["file_path"])
            if not apath:
                st.stop()

        if master_audio_enabled:
            import os
            import sys
            import zipfile
            import shutil
            import urllib.request
            import subprocess
            import streamlit as st

            # --- 1. SETUP: Download and prepare the vocal-remover package ---
            st.markdown("### Step 1: Setting up the vocal-remover environment")

            if not os.path.exists("vocal-remover"):
                url = "https://github.com/tsurumeso/vocal-remover/releases/download/v6.0.0b4/vocal-remover-v6.0.0b4.zip"
                zip_path = "vocal-remover.zip"

                with st.spinner("Downloading vocal-remover release package..."):
                    urllib.request.urlretrieve(url, zip_path)

                with st.spinner("Unzipping the package..."):
                    with zipfile.ZipFile(zip_path, "r") as zip_ref:
                        zip_ref.extractall(".")

                with st.spinner("Configuring directory..."):
                    if os.path.exists("vocal-remover-v6.0.0b4"):
                        shutil.move("vocal-remover-v6.0.0b4", "vocal-remover")

                with st.spinner("Installing dependencies..."):
                    subprocess.check_call([sys.executable, "-m", "pip", "install", "-r", "vocal-remover/requirements.txt"])

                st.success("✅ Setup complete.")
            else:
                st.info("✅ Vocal-remover environment already set up.")

            # --- 2. FIX IMPORTS: Add the vocal-remover library to the Python path ---
            if "vocal-remover" not in sys.path:
                sys.path.append("vocal-remover")
            
            import numpy as np
            import soundfile as sf
            import torch
            from lib import nets
            from lib import spec_utils
            from inference import Separator

            with st.spinner("Mastering audio..."):
                apath = master_audio(apath)
                if not apath:
                    st.stop()

        # CRITICAL: Store audio path in session state for sentiment analysis
        st.session_state["audio_path"] = apath
        st.write(f"Debug: Audio path stored: {apath}")  # Debug line
        
        tr = transcribe_with_whisperx(apath)
        if not tr:
            st.error("Transcription failed."); st.stop()

        # CRITICAL: Store whisper result in session state for sentiment analysis
        st.session_state["whisper_result"] = tr
        st.write(f"Debug: Whisper result stored with {len(tr.get('segments', []))} segments")  # Debug line

        diar_path = vm["base"] + "_diarized_transcript.txt"
        st.session_state.diar_path = diar_path
        

        try:
            diar = diarize_with_whisperx(apath, tr)
            if diar and "segments" in diar:
                format_diarized_transcript(diar, diar_path)
                st.success(f"Diarized transcript saved: {diar_path}")
                # Update whisper result with diarization info if available
                st.session_state["whisper_result"] = diar
                st.write("Debug: Updated whisper result with diarization data")  # Debug line
        except Exception as e:
            st.warning(f"Diarization error, falling back: {e}")
            # fallback: timestamped from segments
            def fmt(seconds: float) -> str:
                if seconds is None: return "??:??:??"
                return str(timedelta(seconds=int(seconds)))
            with open(diar_path, "w", encoding="utf-8") as f:
                for seg in tr.get("segments", []):
                    stt = fmt(seg.get("start")); en = fmt(seg.get("end"))
                    text = (seg.get("text") or "").strip()
                    f.write(f"[{stt} - {en}] SPEAKER_UNKNOWN: {text}\n")
            st.success(f"Timestamped transcript saved: {diar_path}")

        # Final debug check
        st.write("Debug: Session state after transcription:")
        st.write(f"- audio_path exists: {'audio_path' in st.session_state}")
        st.write(f"- whisper_result exists: {'whisper_result' in st.session_state}")
        if 'whisper_result' in st.session_state:
            st.write(f"- segments count: {len(st.session_state['whisper_result'].get('segments', []))}")

if generate_chapters_btn:
    if 'whisper_result' not in st.session_state or not st.session_state['whisper_result']:
        st.error("Please transcribe the video first.")
    else:
        with st.spinner("Generating chapters..."):
            chapters = generate_chapters_with_llama("", st.session_state['whisper_result']['segments'])
            if chapters:
                st.header("Generated Chapters")
                for chap in chapters:
                    st.subheader(f"{chap['title']}")
                    st.write(f"**Time:** {chap['start_time']} - {chap['end_time']}")
                    st.write(f"**Description:** {chap['description']}")
            else:
                st.error("Failed to generate chapters.")

if run_summary_btn:
    diar_path = st.session_state.get("diar_path")
    if not diar_path or not os.path.exists(diar_path):
        st.error("Please run Transcribe + Diarize at least once before summarizing.")
    else:
        base_text = None
        if use_official_yt and st.session_state.yt_transcript_available and st.session_state.yt_transcript_text:
            base_text = st.session_state.yt_transcript_text
            st.info("Using official YouTube transcript for summarization.")
        if base_text is None:
            with open(diar_path, "r", encoding="utf-8") as f:
                base_text = f.read()

        with st.spinner("Summarizing via Nebius map-reduce…"):
            s0 = time.time()
            goal = goal_map[sum_goal_label]
            sres = summarize_long_transcript_map_reduce(base_text, goal, sum_len)
            s1 = time.time()

        if not sres:
            st.error("Summarization failed.")
        else:
            st.subheader("Summary")
            st.write(sres["final_report"])
            st.caption(f"Took {s1-s0:.2f}s")

            # Downloads
            sum_path = vm["base"] + f"_{goal.value}_summary.md"
            with open(sum_path, "w", encoding="utf-8") as f:
                f.write(sres["final_report"])
            st.download_button(
                "Download summary (.md)",
                data=open(sum_path,"rb").read(),
                file_name=os.path.basename(sum_path),
                mime="text/markdown"
            )

            if sres.get("topics"):
                import json
                st.subheader("Extracted Topics & Keywords")
                st.json(sres["topics"])
                topics_path = vm["base"] + "_topics.json"
                with open(topics_path, "w", encoding="utf-8") as f:
                    json.dump(sres["topics"], f, indent=2)
                st.download_button(
                    "Download topics (.json)",
                    data=open(topics_path,"rb").read(),
                    file_name=os.path.basename(topics_path),
                    mime="application/json"
                )
   # ================= NEW: SENTIMENT ANALYSIS UI =================
    # Replace the sentiment analysis section in the sidebar (around line 680-750) with this fixed version:

st.divider()
st.header("5) Sentiment Analysis")

# Check what we have in session state
has_audio = st.session_state.get("audio_path") is not None and os.path.exists(st.session_state.get("audio_path", ""))
has_whisper = st.session_state.get("whisper_result") is not None
has_diarized = st.session_state.get("diar_path") is not None

# Debug information (can be removed once working)
st.write(f"Debug - Has audio: {has_audio}")
st.write(f"Debug - Has whisper: {has_whisper}")
st.write(f"Debug - Has diarized: {has_diarized}")

if has_audio and has_whisper:
    whisper_result = st.session_state["whisper_result"]
    
    # Calculate total duration safely
    total_duration = 0
    if whisper_result and whisper_result.get("segments"):
        try:
            end_times = [seg.get('end', 0) for seg in whisper_result["segments"] if seg.get('end') is not None]
            if end_times:
                total_duration = max(end_times)
                st.write(f"Debug - Calculated duration: {total_duration:.2f} seconds")
        except Exception as e:
            st.warning(f"Could not calculate duration: {e}")
            total_duration = 60  # fallback
    
    # Always show the controls if we have the required data, regardless of duration calculation
    if total_duration <= 0:
        # Fallback: estimate from segment count
        segment_count = len(whisper_result.get("segments", []))
        total_duration = max(60, segment_count * 3)  # rough estimate: 3 seconds per segment
        st.info(f"Using estimated duration: {total_duration:.2f} seconds based on {segment_count} segments")
    
    st.success(f"Sentiment analysis ready! Total duration: {total_duration:.2f} seconds")
    
    # Time range inputs
    col1, col2 = st.columns(2)
    with col1:
        sentiment_start = st.number_input(
            "Start time (seconds)", 
            min_value=0.0, 
            max_value=total_duration, 
            value=0.0, 
            step=1.0,
            key="sentiment_start_time"
        )
    with col2:
        sentiment_end = st.number_input(
            "End time (seconds)", 
            min_value=sentiment_start + 1.0,
            max_value=total_duration, 
            value=min(sentiment_start + 10, total_duration), 
            step=1.0,
            key="sentiment_end_time"
        )
    
    sentiment_type = st.selectbox(
        "Analysis Type", 
        ["Text-Only", "Speech-Only", "Multimodal"],
        key="sentiment_analysis_type"
    )
    
    # Show preview of selected segment
    with st.expander("Preview selected segment"):
        preview_text = ""
        if whisper_result.get("segments"):
            for seg in whisper_result["segments"]:
                seg_start = seg.get('start', 0)
                seg_end = seg.get('end', 0)
                if seg_start < sentiment_end and seg_end > sentiment_start:
                    preview_text += seg.get('text', '') + " "
        
        if preview_text.strip():
            st.write(f"Text in range: {preview_text.strip()[:200]}...")
        else:
            st.write("No text found in selected range")
    
    run_sentiment_btn = st.button("🎭 Run Sentiment Analysis", key="run_sentiment_analysis")
    
else:
    missing_items = []
    if not has_audio:
        missing_items.append("audio file")
    if not has_whisper:
        missing_items.append("transcription result")
    
    st.info(f"Run 'Transcribe + Diarize' first to enable sentiment analysis. Missing: {', '.join(missing_items)}")
    run_sentiment_btn = False
# 4) Build/Load RAG
if build_btn:
    diar_path = st.session_state.get("diar_path")
    if not diar_path or not os.path.exists(diar_path):
        st.error("Please transcribe/diarize first.")
    else:
        with st.spinner("Building/Loading FAISS index..."):
            st.session_state.index_dir = st.session_state.video_meta["base"] + "_faiss_index"
            qna = initialize_sentence_rag(diar_path, st.session_state.index_dir, st.session_state.video_meta)
            if qna:
                st.session_state.qna = qna
                st.success("RAG index ready.")
            else:
                st.error("Failed to init RAG.")


st.divider()
st.subheader("Chat with the transcript")
if "qna" in st.session_state and st.session_state.qna:
    # Display existing chat messages
    for role, msg in st.session_state.chat_history:
        with st.chat_message(role):
            st.markdown(msg)

    transcribed_prompt = None

    # Voice input section
    try:
        from streamlit_mic_recorder import mic_recorder
        import soundfile as sf
        import numpy as np
        
        # Check if all voice dependencies are available
        voice_available = all([mic_recorder, sf, np])
        
    except ImportError:
        voice_available = False

    if voice_available:
        st.write("🎙️ **Voice Input**")
        
        # Configure mic recorder with better settings
        audio_info = mic_recorder(
            start_prompt="🎙️ Click to start recording",
            stop_prompt="🛑 Recording... Click to stop",
            just_once=False,
            use_container_width=True,
            callback=None,
            args=(),
            kwargs={},
            key='voice_recorder'
        )
        
        # Process audio if received
        if audio_info and st.session_state.get('voice_recorder'):
            st.write("Processing your voice input...")
            with st.spinner("Transcribing your question..."):
                transcribed_prompt = transcribe_voice_query(audio_info)
                
            if transcribed_prompt and transcribed_prompt != "Could not understand the audio. Please try again.":
                st.success(f"Voice input: '{transcribed_prompt}'")
            elif transcribed_prompt:
                st.warning(transcribed_prompt)
    else:
        st.info("📦 To enable voice input, install: `pip install streamlit-mic-recorder soundfile numpy`")

    # Text input
    text_prompt = st.chat_input("💬 Or type your question here...")
    
    # Use whichever prompt is available
    prompt = transcribed_prompt if transcribed_prompt else text_prompt

    if prompt and prompt != "Could not understand the audio. Please try again.":
        is_voice_prompt = transcribed_prompt is not None and transcribed_prompt == prompt

        st.session_state.chat_history.append(("user", prompt))
        with st.chat_message("user"):
            if is_voice_prompt:
                st.markdown(f"🎙️ {prompt}")
            else:
                st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                t0 = time.time()
                answer, link = st.session_state.qna.answer(prompt)
                t1 = time.time()
                st.markdown(answer)
                st.caption(f"Latency: {t1 - t0:.2f}s")
                if link and st.session_state.video_meta.get("source") == "youtube":
                    st.markdown(f"[Open segment]({link})")
                elif link and st.session_state.video_meta.get("source") == "file":
                    st.code(link, language="bash")
                    
        st.session_state.chat_history.append(("assistant", answer))

        # Clear the voice recorder state after processing
        if is_voice_prompt:
            st.session_state['voice_recorder'] = None

        st.rerun()

else:
    st.info("Load a video, Transcribe+Diarize, then Build RAG to start chatting.")


if 'run_sentiment_btn' in locals() and run_sentiment_btn:
    if st.session_state.get("audio_path") and st.session_state.get("whisper_result"):
        audio_path = st.session_state["audio_path"]
        whisper_result = st.session_state["whisper_result"]
        
        # Get the values from session state with proper keys
        sentiment_start = st.session_state.get("sentiment_start_time", 0.0)
        sentiment_end = st.session_state.get("sentiment_end_time", 10.0)
        sentiment_type = st.session_state.get("sentiment_analysis_type", "Text-Only")
        
        with st.spinner("Running sentiment analysis..."):
            try:
                result = run_sota_emotion_analysis(
                    audio_path, 
                    whisper_result, 
                    sentiment_start, 
                    sentiment_end, 
                    sentiment_type
                )
                
                if result:
                    st.success("🎭 Sentiment Analysis Complete!")
                    
                    # Create results section
                    st.subheader("Emotion Analysis Results")
                    
                    col1, col2 = st.columns([1, 2])
                    
                    with col1:
                        st.metric("Dominant Emotion", result["dominant"])
                        if result.get("note"):
                            st.info(result["note"])
                    
                    with col2:
                        st.write("**Emotion Scores:**")
                        for emotion, score in result["scores"].items():
                            # Create a progress bar for each emotion
                            progress_value = max(0.0, min(1.0, score))  # Clamp between 0 and 1
                            st.progress(progress_value, text=f"{emotion}: {score:.3f}")
                    
                    # Additional details
                    with st.expander("Analysis Details"):
                        st.write(f"**Time Range:** {sentiment_start:.2f}s - {sentiment_end:.2f}s")
                        st.write(f"**Analysis Type:** {sentiment_type}")
                        st.write(f"**Total Score:** {result['total_score']:.3f}")
                        
                else:
                    st.error("Sentiment analysis failed. Please try a different time range or analysis type.")
            
            except Exception as e:
                st.error(f"Error during sentiment analysis: {str(e)}")
                st.write("Debug info:", str(e))
# ============================ FLASHCARDS ============================
st.divider()
st.subheader("🗂️ Flashcard Generator")
col_f1, col_f2 = st.columns([2,1])
with col_f1:
    flash_source = st.radio("Use content from:", ["Summary (if generated)", "Diarized Transcript"], horizontal=True)
    n_cards = st.slider("Number of flashcards", 5, 50, 20, 1)
run_flash_btn = st.button("Generate Flashcards")

if run_flash_btn:
    diar_path = st.session_state.get("diar_path")
    if not diar_path or not os.path.exists(diar_path):
        st.error("Please run Transcribe + Diarize first.")
    else:
        # prefer last saved summary markdown if present
        base_text = None
        if flash_source == "Summary (if generated)":
            # try to find any _summary.md saved next to base file
            base = st.session_state.video_meta.get("base")
            if base:
                try:
                    cand = [p for p in os.listdir(os.path.dirname(base) or ".")
                            if p.startswith(os.path.basename(base)) and p.endswith("_summary.md")]
                except Exception:
                    cand = []
                if cand:
                    cand_path = os.path.join(os.path.dirname(base) or ".", sorted(cand)[-1])
                    try:
                        base_text = open(cand_path, "r", encoding="utf-8").read()
                    except Exception:
                        base_text = None
        if base_text is None:
            base_text = open(diar_path, "r", encoding="utf-8").read()

        with st.spinner("Generating flashcards…"):
            cards = generate_flashcards_from_text(base_text, n_cards)
        if not cards:
            st.error("Failed to generate flashcards.")
        else:
            st.success(f"Generated {len(cards)} flashcards.")
            for i, c in enumerate(cards, 1):
                with st.expander(f"Card {i}: {c['question'][:80]}"):
                    st.markdown(f"**Q:** {c['question']}\n\n**A:** {c['answer']}")
            # downloads
            import json
            st.download_button("Download JSON", data=json.dumps(cards, indent=2).encode("utf-8"),
                               file_name="flashcards.json", mime="application/json")
            st.download_button("Download CSV", data=flashcards_to_csv(cards),
                               file_name="flashcards.csv", mime="text/csv")
            st.download_button("Download TSV (Anki-friendly)", data=flashcards_to_tsv(cards),
                               file_name="flashcards.tsv", mime="text/tab-separated-values")

# ============================ EXPORTS ===============================
st.divider()
st.subheader("📥 Download Answers (Conversation Export)")
if st.session_state.chat_history:
    md = conversation_to_markdown(st.session_state.chat_history)
    st.download_button("Download Markdown", data=md.encode("utf-8"), file_name="conversation.md", mime="text/markdown")

    docx_bytes = conversation_to_docx_bytes(st.session_state.chat_history)
    if docx_bytes:
        st.download_button("Download DOCX", data=docx_bytes, file_name="conversation.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.info("Install `python-docx` to enable DOCX export: `pip install python-docx`")

    # PDF export: swallow any unexpected errors so UI never shows a traceback
    pdf_bytes = None
    try:
        pdf_bytes = conversation_to_pdf_bytes(st.session_state.chat_history)
    except Exception:
        pdf_bytes = None
    if pdf_bytes:
        st.download_button("Download PDF", data=pdf_bytes, file_name="conversation.pdf", mime="application/pdf")
    else:
        if FPDF is None:
            st.info("Install `fpdf2` to enable PDF export: `pip install fpdf2`")
        else:
            st.warning("PDF export had a rendering hiccup. Try Markdown/DOCX or remove extremely long unbroken text.")
else:
    st.info("Ask a question first to enable conversation export.")